# -*- coding: utf-8 -*-
"""生成《系统部署说明》与《测试数据》两份 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

GREEN = RGBColor(0x1A, 0x70, 0x51)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x59, 0x59, 0x59)


def set_font(run, size=12, bold=False, color=DARK, name='宋体', western='Times New Roman'):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = western
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_para(doc, text, size=12, bold=False, color=DARK, name='宋体',
             space_after=6, align=None, line=1.4, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, name=name)
    return p


def add_title(doc, text, sub=None):
    add_para(doc, text, size=18, bold=True, color=GREEN, name='黑体',
             space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    if sub:
        add_para(doc, sub, size=12, color=GRAY, name='楷体', space_after=16,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_para(doc, '', space_after=10)


def add_h2(doc, text):
    add_para(doc, text, size=14, bold=True, color=GREEN, name='黑体',
             space_after=8)


def add_h3(doc, text):
    add_para(doc, text, size=12.5, bold=True, color=DARK, name='黑体',
             space_after=6)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.left_indent = Pt(21)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x2A, 0x4A, 0x3A)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '1A7051')
        hdr[i]._tc.get_or_add_tcPr().append(shd)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(v))
            set_font(run, size=10.5, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for r in table.rows:
                r.cells[i].width = Cm(w)
    add_para(doc, '', space_after=8)
    return table


def setup_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    return doc


# ============================================================
# 文档一：系统部署说明
# ============================================================
doc = setup_doc()
add_title(doc, '系统部署说明', '—— 学途智面·面试罗盘（岗位能力图谱求职平台）')

add_h2(doc, '一、系统架构概述')
add_para(doc, '本系统采用前后端分离的微服务化架构，由五个服务组件构成：PostgreSQL 关系型数据库、Neo4j 图数据库（Docker 容器化部署）、FastAPI 知识图谱服务、Flask 业务后端以及 Vite 前端应用。各组件说明如下表所示。')
add_table(doc,
    ['组件', '技术栈', '端口', '职责'],
    [
        ['数据库', 'PostgreSQL', '5432', '存储用户、岗位、资源、面试记录等业务数据'],
        ['图数据库', 'Neo4j（Docker）', '7474 / 7687', '存储岗位-技能知识图谱，支持 Cypher 图查询'],
        ['图谱服务', 'Python FastAPI', '7576', '图谱构建、人岗匹配、趋势分析、相似岗位推荐'],
        ['业务后端', 'Python Flask', '8082', '认证、岗位、问答、资源、面试等业务接口'],
        ['前端', 'React + Vite', '5173', '页面渲染，通过代理转发 /api 请求至后端'],
    ],
    widths=[2.5, 3.5, 2.5, 7.5])

add_h2(doc, '二、运行环境要求')
add_table(doc,
    ['环境项', '版本/要求', '说明'],
    [
        ['操作系统', 'Windows 10/11（64 位）', 'Linux 亦可，启动命令相应调整'],
        ['Python', '3.10 及以上', '需安装 flask、fastapi、uvicorn、neo4j、psycopg2 等'],
        ['Node.js', '18 及以上', '前端构建与本地运行'],
        ['Docker Desktop', '最新稳定版', '用于运行 Neo4j 容器'],
        ['PostgreSQL', '16 及以上', '本机安装于 D:\\pgsql'],
        ['外部服务', '阿里云百炼 / Coze 扣子', '大模型与智能体工作流，需有效 API Key'],
    ],
    widths=[3, 5, 8])

add_h2(doc, '三、部署步骤')

add_h3(doc, '步骤 1：启动 PostgreSQL 数据库')
add_para(doc, '若系统服务方式启动失败（权限问题），可采用 pg_ctl 直接启动，命令如下：')
add_code(doc, 'D:\\pgsql\\bin\\pg_ctl.exe start -D "D:\\pgsql\\data"')
add_para(doc, '验证：执行 netstat -ano | findstr 5432，出现 LISTENING 即为成功。数据库默认连接信息：地址 127.0.0.1，端口 5432，用户 postgres，密码 123456，库名 job_graph。')

add_h3(doc, '步骤 2：初始化数据库')
add_para(doc, '按以下顺序在 job_graph 库中执行 SQL 脚本（脚本位于 scripts/data_import 目录）：')
add_table(doc,
    ['顺序', '脚本', '作用'],
    [
        ['1', 'init_all_schema.sql', '创建全部业务表结构与索引'],
        ['2', 'data.sql', '初始化 4 个角色与 3 个测试用户'],
        ['3', 'learning-resources-schema.sql', '创建学习资源与聊天记录相关表'],
        ['4', 'learning-resources-data.sql', '导入 12 个资源系列、401 篇文档'],
        ['5', 'insert_jobs.sql', '写入 7 条示例岗位'],
        ['6', 'import_job_data_fixed.sql', '导入 165 条 Boss 直聘真实岗位（会清空并覆盖旧岗位）'],
    ],
    widths=[1.5, 5.5, 9])

add_h3(doc, '步骤 3：启动 Neo4j 图数据库')
add_para(doc, '在项目根目录执行（docker-compose-neo4j-only.yml）：')
add_code(doc, 'docker compose -f docker-compose-neo4j-only.yml up -d')
add_para(doc, '验证：浏览器访问 http://localhost:7474 能打开管理界面；Bolt 端口 7687 供服务连接。默认账号 neo4j，密码 password123。注：容器健康检查偶显 unhealthy，但 7474/7687 可访问即功能正常。')

add_h3(doc, '步骤 4：配置环境变量')
add_para(doc, '在 backend 目录创建 .env 文件（可参照 .env.example），内容如下：')
add_code(doc,
'''# 数据库连接配置
DATABASE_URL=postgresql://postgres:123456@localhost:5432/job_graph

# 阿里云百炼（大模型）
DASHSCOPE_API_KEY=sk-xxxxxxxxxxxxxxxx
DASHSCOPE_BASE_URL=https://dashscope.aliyuncs.com/compatible-mode/v1
LLM_MODEL=qwen-plus
LLM_SMALL_MODEL=qwen-turbo
LLM_EMBEDDING_MODEL=text-embedding-v3

# Neo4j 连接配置
NEO4J_URI=bolt://localhost:7687
NEO4J_USER=neo4j
NEO4J_PASSWORD=password123

# Coze（扣子）智能体工作流
COZE_API_TOKEN=pat_xxxxxxxxxxxxxxxx
COZE_WORKFLOW_ID=xxxxxxxxxxxxxxxx
COZE_BASE_URL=https://api.coze.cn''')
add_para(doc, '注意：.env 含密钥，已被 .gitignore 排除，切勿提交仓库或外传。')

add_h3(doc, '步骤 5：安装依赖并启动图谱服务（7576）')
add_code(doc, 'cd backend\npip install fastapi uvicorn neo4j psycopg2-binary httpx requests python-dotenv pydantic\npython job_skill_graph_service.py')

add_h3(doc, '步骤 6：启动 Flask 业务后端（8082）')
add_code(doc, 'cd backend\npython full_app.py')

add_h3(doc, '步骤 7：启动前端（5173）')
add_code(doc, 'cd frontend\nnpm install\nnpm run dev')
add_para(doc, '前端开发服务器自带代理：/api 转发至 Flask（8082），/api/job-skill-graph 转发至 FastAPI（7576），无需额外配置跨域。')

add_h2(doc, '四、部署验证')
add_table(doc,
    ['检查项', '验证方式', '预期结果'],
    [
        ['PostgreSQL', 'netstat -ano | findstr 5432', 'LISTENING'],
        ['Neo4j', '访问 http://localhost:7474', '出现管理界面登录页'],
        ['图谱服务', 'GET http://127.0.0.1:7576/docs', 'HTTP 200，Swagger 文档可见'],
        ['业务后端', 'POST http://127.0.0.1:8082/api/auth/login', 'HTTP 200，返回 token'],
        ['前端', '访问 http://localhost:5173', '首页正常渲染'],
    ],
    widths=[3, 7, 6])

add_h2(doc, '五、默认测试账号')
add_table(doc,
    ['账号', '密码', '角色', '说明'],
    [
        ['student1', '123456', '学生', '真实数据库账号'],
        ['company1', '123456', '企业', '真实数据库账号'],
        ['admin', '123456', '管理员', '真实数据库账号'],
        ['student / company / school', '123456', '三角色', '前端 Mock 演示账号'],
    ],
    widths=[4, 2.5, 3, 6.5])

add_h2(doc, '六、常见问题排查')
add_table(doc,
    ['现象', '原因', '解决办法'],
    [
        ['后端启动报数据库连接失败', 'PostgreSQL 未启动', '按步骤 1 启动后重启后端'],
        ['图谱服务报 DASHSCOPE_API_KEY 未配置', '.env 缺失或未加载', '检查 backend/.env 并重启服务'],
        ['LLM 调用返回 400 Arrearage', '百炼账户欠费', '登录阿里云百炼控制台充值'],
        ['接口 404', '前端请求路径与后端路由不一致', '核对 /api 前缀与后端 @app.route 定义'],
        ['端口被占用', '旧进程未退出', 'netstat 查 PID 后 taskkill /PID xxx /F'],
        ['图谱数据为空', '图谱未构建', '调用 /api/job-skill-graph/build-graph 触发构建'],
    ],
    widths=[4.5, 4.5, 7])

doc.save(r'd:\jianli\resume_analysis_system\系统部署说明.docx')
print('saved: 系统部署说明.docx')

# ============================================================
# 文档二：测试数据
# ============================================================
doc2 = setup_doc()
add_title(doc2, '系统测试数据说明', '—— 学途智面·面试罗盘（岗位能力图谱求职平台）')

add_h2(doc2, '一、测试数据总体统计')
add_table(doc2,
    ['数据类别', '数量', '存储位置', '来源'],
    [
        ['岗位数据', '165 条', 'job_postings 表', 'Boss 直聘真实爬取'],
        ['示例岗位', '7 条（被上表覆盖）', 'job_postings 表', '手工构造'],
        ['角色数据', '4 条', 'roles 表', '学生/学校/企业/管理员'],
        ['用户数据', '3 条', 'users 表', 'student1/company1/admin'],
        ['学习资源系列', '12 个', 'resource_series 表', 'MDN Web 文档等'],
        ['学习文档', '401 篇', 'documents 表', '结构化技术文档'],
        ['图谱岗位节点', '49 个', 'Neo4j Job 节点', '图谱构建流水线'],
        ['图谱技能节点', '213 个', 'Neo4j Skill 节点', '大模型技能规范化'],
        ['REQUIRES 关系', '367 条', 'Neo4j 关系', '带 required/preferred 属性'],
    ],
    widths=[3.5, 3.5, 4, 5])

add_h2(doc2, '二、测试账号数据')
add_table(doc2,
    ['用例编号', '账号', '密码', '角色', '预期结果'],
    [
        ['TC-AUTH-01', 'student1', '123456', '学生', '登录成功，返回 token'],
        ['TC-AUTH-02', 'company1', '123456', '企业', '登录成功，进入企业端菜单'],
        ['TC-AUTH-03', 'admin', '123456', '管理员', '登录成功，可进入图谱运维后台'],
        ['TC-AUTH-04', 'student1', 'wrong123', '学生', '登录失败，提示账号或密码错误'],
        ['TC-AUTH-05', 'notexist', '123456', '无', '登录失败，提示账号或密码错误'],
    ],
    widths=[2.8, 3, 2.5, 2.5, 5.2])

add_h2(doc2, '三、注册模块测试数据')
add_table(doc2,
    ['用例编号', '测试项', '输入数据', '预期结果'],
    [
        ['TC-REG-01', '正常注册', '用户名 teststu01 / 真实姓名 张测试 / 邮箱 test01@qq.com / 手机 13800138001 / 密码 Test@1234', '注册成功，跳转登录'],
        ['TC-REG-02', '用户名过短', '用户名 ab', '提示"用户名至少3位"'],
        ['TC-REG-03', '用户名含非法字符', '用户名 stu@01', '提示仅允许字母数字下划线'],
        ['TC-REG-04', '邮箱格式错误', '邮箱 abc@@qq', '提示"请输入有效的邮箱地址"'],
        ['TC-REG-05', '手机号错误', '手机 12345', '提示"请输入有效的手机号码"'],
        ['TC-REG-06', '两次密码不一致', '密码 Test@1234 / 确认 Test@4321', '提示两次输入不一致'],
    ],
    widths=[2.6, 3, 6, 4.4])

add_h2(doc2, '四、求职市场模块测试数据')
add_table(doc2,
    ['用例编号', '测试项', '输入数据', '预期结果'],
    [
        ['TC-JOB-01', '关键词搜索', '关键词：python', '返回岗位名称含 python 的岗位列表'],
        ['TC-JOB-02', '关键词搜索', '关键词：Java', '返回 Java 相关岗位'],
        ['TC-JOB-03', '空结果搜索', '关键词：zzzzz', '返回空列表，页面显示暂无数据'],
        ['TC-JOB-04', '条件筛选', '工作类型：全职；薪资范围筛选', '列表按条件过滤'],
        ['TC-JOB-05', '岗位详情', '点击任一岗位卡片', '展示技能要求、薪资、福利等完整字段'],
        ['TC-JOB-06', '首次投递', 'student1 对目标岗位投递', '投递成功'],
        ['TC-JOB-07', '重复投递校验', '同一账号 7 天内再次投递同一岗位', '系统拦截，提示不可重复投递'],
    ],
    widths=[2.6, 3, 6, 4.4])

add_h2(doc2, '五、人岗匹配模块测试数据')
add_para(doc2, '标准测试简历文本（复制使用）：')
add_code(doc, 'ggq，24岁，刚毕业，想从事python相关行业，在校期间有拿过很多相关比赛奖项。熟悉Python、MySQL，了解Django框架。')
add_table(doc2,
    ['用例编号', '目标岗位', '简历输入', '预期结果'],
    [
        ['TC-MATCH-01', 'Python开发', '上述标准简历', '匹配成功，输出综合分、五维雷达图、已具备技能含 Python/MySQL，缺失技能清单非空'],
        ['TC-MATCH-02', 'Python开发', '空文本', '前端提示请填写简历后分析'],
        ['TC-MATCH-03', '不存在的岗位', '上述标准简历', '系统提示未匹配到岗位或给出近似推荐'],
        ['TC-MATCH-04', 'Python开发', '长文本简历（800字以上）', '正常解析，响应时间可接受'],
        ['TC-MATCH-05', 'Python开发', '上述标准简历（LLM 服务不可用）', '自动降级关键词匹配，仍输出完整评分结果'],
    ],
    widths=[2.6, 2.8, 5.6, 5])

add_h2(doc2, '六、简历分析模块测试数据')
add_para(doc2, '标准测试岗位要求（复制使用）：')
add_code(doc, '本科及以上学历，计算机相关专业；3年以上Python开发经验；熟悉 Django/Flask 框架；熟悉 MySQL、Redis；有分布式系统经验者优先。')
add_para(doc2, '标准测试简历文本（复制使用）：')
add_code(doc, '张三，本科，计算机科学与技术专业，3年Python后端开发经验。熟练使用 Django、Flask 框架，熟悉 MySQL、Redis，参与过电商订单系统开发，了解 Docker 部署。')
add_table(doc2,
    ['用例编号', '岗位要求', '简历文本', '预期结果'],
    [
        ['TC-RES-01', '上述岗位要求', '上述简历', '输出五字段 JSON：岗位匹配度/综合评估/亮点/待改进/面试建议，且亮点与简历内容对应'],
        ['TC-RES-02', '上述岗位要求', '空白', '提示请填写完整信息'],
        ['TC-RES-03', '空白', '上述简历', '提示请填写完整信息'],
        ['TC-RES-04', '上述岗位要求', '只有姓名年龄的极简简历', '正常输出，待改进项包含技能与经验缺失'],
    ],
    widths=[2.6, 3.4, 4.6, 5.4])

add_h2(doc2, '七、求职问答模块测试数据')
add_table(doc2,
    ['用例编号', '输入问题', '预期结果'],
    [
        ['TC-QA-01', 'Java 开发方向需要掌握哪些技能？', '返回与 Java 技能栈相关的结构化建议，非固定模板'],
        ['TC-QA-02', '如何准备校园招聘面试？', '返回分步骤的面试准备建议'],
        ['TC-QA-03', '（空消息）', '前端拦截，不发送请求'],
        ['TC-QA-04', '超长问题（500字）', '正常返回，不报错'],
        ['TC-QA-05', 'Coze 服务不可用时发送任意问题', '自动降级 LLM 链路，仍返回有效回答'],
    ],
    widths=[2.6, 6, 7.4])

add_h2(doc2, '八、能力图谱与趋势分析测试数据')
add_table(doc2,
    ['用例编号', '操作', '预期结果'],
    [
        ['TC-GRAPH-01', '选择岗位"Python开发"查看图谱', '展示该岗位核心/优先技能节点与关系'],
        ['TC-GRAPH-02', '点击技能节点（如 Python）', '反查关联岗位列表与共现技能'],
        ['TC-GRAPH-03', '查看图谱统计', '节点/关系数量与 Neo4j 实际一致（49/213/367 量级）'],
        ['TC-TREND-01', '打开趋势分析页', '展示热门技能 Top15、技能共现关系、新兴岗位三类视图'],
        ['TC-TREND-02', '查看 AI 解读', '基于统计数据生成自然语言解读'],
        ['TC-TREND-03', 'LLM 不可用时查看趋势', '统计图表正常展示，解读区提示未配置'],
    ],
    widths=[2.8, 5.6, 7.6])

add_h2(doc2, '九、模拟面试与学习资源测试数据')
add_table(doc2,
    ['用例编号', '操作', '预期结果'],
    [
        ['TC-IV-01', 'student1 开始一场 Python 岗模拟面试并作答', '生成面试报告，含总分与分项得分'],
        ['TC-IV-02', '查看历史面试记录列表', '按时间倒序展示，可进入详情与回放'],
        ['TC-RES-01', '浏览学习资源', '展示 12 个资源系列、401 篇文档'],
        ['TC-RES-02', '打开任一文档', '正文正常渲染，图片资源可加载'],
        ['TC-RES-03', '收藏/取消收藏文档', '收藏状态即时更新并可查询'],
        ['TC-RES-04', '查看智能推荐', '推荐内容与人岗匹配诊断出的缺失技能相关'],
    ],
    widths=[2.8, 6.2, 7])

doc2.save(r'd:\jianli\resume_analysis_system\测试数据.docx')
print('saved: 测试数据.docx')
