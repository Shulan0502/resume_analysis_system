# -*- coding: utf-8 -*-
"""生成功能介绍 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

GREEN = RGBColor(0x1A, 0x70, 0x51)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x59, 0x59, 0x59)

FEATURES = [
    ("5.1.1 账号注册与登录功能（图 5-1）",
     "支持学生、学校、企业三种角色注册登录，前后端双重校验用户名、邮箱、手机号与密码强度；登录后按角色呈现差异化菜单，超级管理员另设独立入口进入图谱运维后台。",
     ["Flask 构建认证接口，密码 MD5 加盐哈希后存入 PostgreSQL，杜绝明文存储",
      "登录后签发 Token 令牌，后续请求凭 Token 鉴权，接口无状态",
      "基于 RBAC 模型实现三角色权限隔离"],
     "认证、授权与业务解耦，便于水平扩展；三角色体系为学校端就业分析与企业端人才筛选预留扩展空间；表单实时校验显著降低脏数据入库概率。"),
    ("5.1.2 求职问答功能（图 5-2）",
     "提供对话式职业咨询，解答岗位技能要求、职业规划、面试准备等问题，支持历史对话查询与删除。",
     ["主链路：调用扣子（Coze）平台智能体工作流，解析返回结果",
      "备用链路：主链路异常时自动降级，直连阿里云百炼 qwen-plus，以 System Prompt 约束回答风格"],
     "双链路互为容错，任一外部服务故障不影响核心体验；智能体工作流支持可视化编排，无需改码即可迭代应答策略；回答锚定系统内真实岗位数据，降低大模型“幻觉”误导风险。"),
    ("5.1.3 求职市场功能（图 5-3）",
     "浏览企业发布岗位，支持关键词搜索、工作类型与薪资筛选，查看详情并一键投递；系统自动校验 7 天内不可重复投递同一岗位。",
     ["接入 Boss 直聘采集的 165 条真实岗位数据，字段覆盖技能、经验、学历、薪资、福利",
      "企业新发布岗位自动进入多源异构数据流水线，驱动图谱增量更新",
      "PostgreSQL 分页查询 + 多条件组合过滤，高频字段建立索引"],
     "真实数据驱动保证岗位信息可信度；“新岗位→图谱更新”自动化链路响应了赛题对数据动态性的要求；索引优化使数据量增长后检索性能依然平稳。"),
    ("5.1.4 能力图谱功能（图 5-4）",
     "选择目标岗位即可查看可视化技能图谱，核心技能与优先技能以节点连线呈现，界面显示当前图谱版本号，支持图谱动态演化管理。",
     ["Neo4j 存储 Job、Skill 节点及带 importance 属性的 REQUIRES 关系",
      "Cypher 语句实现邻居查询、关系遍历与子图裁剪",
      "前端 ECharts 力导向图渲染，支持节点拖拽、缩放与点击下钻"],
     "图数据库将多层关联查询复杂度降至常数级，天然适配“岗位—技能—共现”深度遍历场景；力导向布局将密集网状数据转化为可读性强的视觉结构，显著降低理解成本。"),
    ("5.1.5 人岗匹配功能（图 5-5）",
     "输入目标岗位与简历文本，系统输出综合匹配分、五维雷达图、已具备技能与待补齐技能清单，实现细粒度人岗差距诊断。",
     ["第一步：大模型从非结构化简历抽取技能、年限、学历等结构化字段",
      "第二步：抽取结果与 Neo4j 图谱能力要求比对，区分必需/优先技能命中情况",
      "第三步：按技能、经验、学历加权算分，环形图+柱状图+雷达图呈现",
      "降级链路：大模型不可用时自动切换图谱技能库关键词匹配与正则抽取"],
     "大模型保证文本解析泛化能力，图谱比对确保匹配口径与真实招聘需求一致，加权评分使结果可量化、可解释，为学习路径规划提供精准输入。"),
    ("5.1.6 趋势分析功能（图 5-6）",
     "提供热门技能排行、技能共现关系、新兴岗位分析三类视图，支持不同时间版本图谱对比，直观呈现技能需求的新增、消退与变迁。",
     ["三类统计均由 Cypher 聚合查询在 Neo4j 侧完成（共现频次、技能重叠度算法等）",
      "统计结果之上由大模型生成自然语言趋势解读报告",
      "新兴岗位通过技能重叠度算法识别，辅助新岗位发现与定义"],
     "以实时图计算替代人工盘点，克服传统手段的“时滞”问题，使“技术爆发→岗位新技能需求”的动态关联可直接查询与验证；AI 解读让统计数据“会说话”。"),
    ("5.1.7 简历分析功能（图 5-7）",
     "提交岗位要求与简历后，从岗位匹配度、综合评估、个人亮点、待改进项、面试建议五个维度生成深度诊断报告。",
     ["调用阿里云百炼 qwen-plus，通过提示词约束严格按 JSON 格式输出五字段",
      "前端多层解析后以卡片形式渲染",
      "外部服务异常时返回基于模板的兜底报告"],
     "结构化输出约束使报告可被程序直接消费；评估结论锚定用户真实输入，有效抑制大模型“幻觉”；五维度报告兼顾能力盘点与行动指引，自然衔接学习资源推荐。"),
    ("5.1.8 模拟面试功能（图 5-8）",
     "提供 AI 仿真面试练习，按目标岗位依次提问、支持追问，全程音视频录制，还原真实面试交互节奏。",
     ["面试问题基于图谱中该岗位的必需与优先技能动态生成",
      "浏览器端 MediaRecorder API 采集音视频，上传服务端存储",
      "录制文件与作答内容关联保存，支持事后回放复盘"],
     "以岗位能力图谱为题目来源，避免通用题库“千人一面”；多模态留存形成“练习—回看—改进”闭环；AI 面试官不受时空限制，显著降低模拟面试边际成本。"),
    ("5.1.9 面试报告功能（图 5-9）",
     "对每次模拟面试量化评估，生成含总分、分项得分与改进建议的报告，支持按时间回溯与音视频回放。",
     ["评估结果写入 PostgreSQL 面试记录表，查询接口按用户维度隔离数据",
      "前端图表渲染分数分布与历史趋势，呈现能力成长曲线"],
     "将主观面试表现转化为可量化、可追踪的数据资产，通过多次报告纵向对比验证学习路径有效性；标准化评分维度也为学校端掌握学生准备度、企业端预判候选人水平提供数据基础。"),
    ("5.1.10 学习资源功能（图 5-10）",
     "内置 400 余篇结构化技术文档（MDN Web 文档、前端框架指南等系列），支持系列浏览、检索、收藏与进度追踪，并依据技能差距智能推荐学习内容。",
     ["资源以“系列—文档”两级结构存储于 PostgreSQL",
      "检索基于模糊匹配与分类过滤实现",
      "推荐逻辑将匹配报告中的缺失技能映射至资源标签，按相关度排序输出"],
     "将抽象技能差距转化为具体可执行的学习内容，解决“只诊断、无处方”问题；文本资源体量轻、更新快、零成本，符合低成本可持续运营定位；进度追踪保障学习连贯性，使能力提升建议真正形成行动闭环。"),
]


def set_font(run, size=12, bold=False, color=DARK, name='宋体'):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_para(doc, text, size=12, bold=False, color=DARK, name='宋体',
             space_after=6, first_indent=0, align=None, line=1.4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(first_indent)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, name=name)
    return p


doc = Document()

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# 文档标题
add_para(doc, '系统功能介绍', size=18, bold=True, color=GREEN, name='黑体',
         space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '—— 基于多源数据与大模型的人才能力图谱平台', size=12,
         color=GRAY, name='楷体', space_after=16, align=WD_ALIGN_PARAGRAPH.CENTER)

for title, brief, techs, adv in FEATURES:
    # 功能标题
    add_para(doc, title, size=14, bold=True, color=GREEN, name='黑体', space_after=8)

    # 功能简介
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    r1 = p.add_run('功能简介：')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(brief)
    set_font(r2, size=12)

    # 技术实现
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run('技术实现：')
    set_font(r, size=12, bold=True)
    for t in techs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.left_indent = Pt(21)
        r = p.add_run('● ' + t)
        set_font(r, size=12, color=GRAY)

    # 技术优势
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.4
    r1 = p.add_run('技术优势：')
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(adv)
    set_font(r2, size=12)

out = r'd:\jianli\resume_analysis_system\系统功能介绍.docx'
doc.save(out)
print('saved:', out)
